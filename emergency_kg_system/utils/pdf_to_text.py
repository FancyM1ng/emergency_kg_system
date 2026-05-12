"""
原始文件 → TXT 统一转换器
===========================
自动识别 data/raw/ 下所有 PDF / HTML / TXT 文件，
提取纯文本到 data/processed/，供三元组抽取流水线使用。

用法:
  python -m utils.pdf_to_text                  # 转换所有支持的文件
  python -m utils.pdf_to_text --dir 企业预案    # 只转换指定子目录
  python -m utils.pdf_to_text --force          # 强制重新转换
"""
import sys
import re
import argparse
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent
RAW_DIR = BASE_DIR / "data" / "raw"
PROCESSED_DIR = BASE_DIR / "data" / "processed"


def pdf_to_text(filepath: Path) -> str:
    """PDF → 纯文本"""
    import pdfplumber
    lines = []
    with pdfplumber.open(str(filepath)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
    return "\n".join(lines)


def html_to_text(filepath: Path) -> str:
    """HTML → 纯文本（保留段落结构）"""
    from bs4 import BeautifulSoup
    with open(filepath, "r", encoding="utf-8") as f:
        html = f.read()

    soup = BeautifulSoup(html, "html.parser")

    # 提取标题
    title_tag = soup.find("title")
    title = title_tag.get_text(strip=True) if title_tag else filepath.stem

    # 移除脚本/样式
    for tag in soup.find_all(["script", "style", "noscript"]):
        tag.decompose()

    body = soup.find("body") or soup

    lines = [title, "=" * min(len(title), 60), ""]
    for element in body.descendants:
        if element.name in ("h1", "h2", "h3", "h4"):
            text = element.get_text(strip=True)
            if text:
                lines.extend(["", text, ""])
        elif element.name == "p":
            text = element.get_text(strip=True)
            if text and len(text) > 5:
                lines.append(text)
        elif element.name == "li":
            text = element.get_text(strip=True)
            if text and len(text) > 2:
                lines.append(f"  - {text}")
        elif element.name in ("td", "th"):
            text = element.get_text(strip=True)
            if text:
                lines.append(text)

    content = "\n".join(lines)
    content = re.sub(r"\n{4,}", "\n\n\n", content).strip()

    # 回退：段落提取太少时用 get_text
    if len(content) < 200:
        plain = body.get_text(separator="\n", strip=True)
        plain = re.sub(r"\n{3,}", "\n\n", plain)
        if len(plain) > len(content):
            content = f"{title}\n{'=' * min(len(title), 60)}\n\n{plain}"

    return content


def docx_to_text(filepath: Path) -> str:
    """DOCX → 纯文本"""
    import docx
    doc = docx.Document(str(filepath))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            if style and style.startswith("Heading"):
                lines.extend(["", text, ""])
            else:
                lines.append(text)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def txt_copy(filepath: Path) -> str:
    """TXT 文件直接读取"""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


CONVERTERS = {
    ".pdf":  pdf_to_text,
    ".html": html_to_text,
    ".htm":  html_to_text,
    ".docx": docx_to_text,
    ".doc":  docx_to_text,
    ".txt":  txt_copy,
}


def get_output_path(raw_file: Path) -> Path:
    """根据 raw 文件路径生成对应的 TXT 输出路径"""
    relative = raw_file.relative_to(RAW_DIR)
    flat_name = str(relative).replace("/", "_").replace("\\", "_")
    return PROCESSED_DIR / Path(flat_name).with_suffix(".txt")


def convert_all(subdir: str = None, force: bool = False):
    """转换所有支持的原始文件"""
    search_dir = RAW_DIR / subdir if subdir else RAW_DIR
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)

    # 收集所有支持格式的文件
    all_files = []
    for ext in CONVERTERS:
        all_files.extend(sorted(search_dir.glob(f"**/*{ext}")))
        all_files.extend(sorted(search_dir.glob(f"**/*{ext.upper()}")))

    if not all_files:
        print(f"No supported files found in: {search_dir}")
        return

    print(f"Found {len(all_files)} files to process\n")

    converted = 0
    skipped = 0
    failed = 0

    for filepath in all_files:
        # 跳过特殊文件
        if filepath.name.startswith("_") or filepath.name.startswith("."):
            continue
        if filepath.suffix.lower() == ".json":
            continue

        output_path = get_output_path(filepath)
        rel_name = str(filepath.relative_to(RAW_DIR))

        if output_path.exists() and not force:
            skipped += 1
            continue

        ext = filepath.suffix.lower()
        converter = CONVERTERS.get(ext)
        if not converter:
            continue

        try:
            print(f"  [{ext}] {rel_name}")
            text = converter(filepath)
            if not text or len(text.strip()) < 50:
                print(f"    WARNING: extracted text too short ({len(text)} chars)")
                failed += 1
                continue

            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
            converted += 1

        except Exception as e:
            print(f"    ERROR: {e}")
            failed += 1

    print(f"\nDone: {converted} converted, {skipped} skipped, {failed} failed")


def main():
    parser = argparse.ArgumentParser(description="Convert raw files to TXT")
    parser.add_argument("--dir", type=str, help="Target subdirectory under data/raw")
    parser.add_argument("--force", action="store_true", help="Force re-conversion")
    args = parser.parse_args()
    convert_all(args.dir, args.force)


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    main()
